from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
def create_cv_docx(llm_output, file_name):
    doc = Document()

    # ---------- Global font ----------
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # ---------- Margins ----------
    section = doc.sections[0]
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

    # ---------- Helpers ----------

    def para_with_date(left_text, right_text, bold=False, size=10.5, space_before=0, space_after=3):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = 1.1
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        # Set a right-aligned tab stop at the right margin
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin,
                            WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

        run = p.add_run(left_text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.bold = bold

        p.add_run('\t')  # This jumps to the right tab stop

        run_date = p.add_run(right_text)
        run_date.font.name = 'Arial'
        run_date.font.size = Pt(size)
        run_date.bold = bold

        return p

    def para(text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=3):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = 1.1
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        run.bold = bold
        return p

    def bullet():
        p = doc.add_paragraph(style='List Bullet')
        pf = p.paragraph_format
        pf.left_indent = Pt(12)
        pf.first_line_indent = Pt(-6)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.1
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        return p

    def add_hyperlink(paragraph, text, url, size=11):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_pr.append(r_fonts)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        r_pr.append(sz)

        run.append(r_pr)
        run.text = text
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    # ---------- Header ----------
    para("Chaitanya Srikanth", bold=True, size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    contact = para("", size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    contact.add_run("E: schaitanya0407@gmail.com | M: +44 7733056417 | ")
    add_hyperlink(contact, "LinkedIn", "https://linkedin.com/in/chaitanyasrikanth", size=11)
    contact.add_run(" | ")
    add_hyperlink(contact, "Git", "https://github.com/codemaster0407", size=11)
    contact.add_run(" | Coventry, UK")

    # ---------- EDUCATION ----------
    para("EDUCATION & QUALIFICATIONS", bold=True, space_before=6, space_after=6)
    para_with_date(
        "MSc in Business Analytics | Warwick Business School, UK",
        "Sep 2025 – Present",
        bold=True
    )
    b = bullet()
    b.add_run(
        "Modules: Business Statistics, Predictive and Prescriptive Analytics, "
        "Marketing Analytics and Technology, Analytics in Practice, Optimisation Models"
    )
    

    para_with_date(
        "Bachelors of Technology in Artificial Intelligence | Mahindra University, India",
        "Sep 2020 – Aug 2024",
        bold=True
    )
    b = bullet()
    b.add_run(
        "Modules: Deep Learning, Machine Learning, Natural Language Processing, Digital Image Processing"
    )
    b = bullet()
    b.add_run("GPA: 8.33 / 10")

    b = bullet()
    b.add_run(
        "Merit Scholarship for academic excellence, "
            "awarded for ranking in the top 10% of the cohort for two consecutive years (2021,2022)"
    )

    # ---------- WORK EXPERIENCE ----------
    para("WORK & LEADERSHIP EXPERIENCE", bold=True, space_before=6, space_after=6)
    para_with_date(
        "Associate AI Engineer | Techolution, India",
        "Jun 2024 – Jul 2025",
        bold=True
    )   

    for pt in llm_output.get("full_time_experience_points", []):
        b = bullet()
        b.add_run(pt)

    para_with_date(
        "AI Intern| Techolution, India",
        "Jun 2024 – Jul 2025",
        bold=True
    )

    for pt in llm_output.get("internship_experience_points", []):
        b = bullet()
        b.add_run(pt)

    # ---------- MENTORING ----------
    para("EXTRA-CURRICULAR EXPERIENCE", bold=True, space_before=8, space_after=6)
    para_with_date(
        "Mentor | Warwick Coding Society, UK",
        "Jun 2024 – Jul 2025",
        bold=True
    )
    mentor_experience = llm_output.get('mentoring_experience', None)
    if mentor_experience == None:
        b = bullet()
        b.add_run("Guiding undergraduate students in Python backend and computer vision application development")
    else:
        for pt in mentor_experience:
            b = bullet()
            b.add_run(pt)

    # ---------- SKILLS ----------
    para("SKILLS AND INTERESTS", bold=True, space_before=8, space_after=6)

    b = bullet()
    b.add_run("Skills: ").bold = True
    b.add_run(llm_output.get("skills") and ", ".join(llm_output["skills"]) or "")

    b = bullet()
    b.add_run("Databases: ").bold = True
    b.add_run(llm_output.get("databases") and ", ".join(llm_output["databases"]) or "")
    b = bullet()
    b.add_run("Cloud: ").bold = True
    b.add_run("AWS, Google Cloud Platform")

    para("ACHIEVEMENTS AND CERTIFICATIONS", bold=True, space_before=8, space_after=4)
    ach_points = llm_output.get('achievements', None)

    b = bullet()
    b.add_run("Certification: ").bold = True
    add_hyperlink(
        b,
        "Google Cloud Certified Professional Machine Learning Engineer",
        "https://www.credly.com/badges/b9a3e1c2-30ea-4b33-8f3b-0755dbe17d5e/linked_in_profile"
    )

    if ach_points == None:
        b = bullet()
        b.add_run(
            "Winner – 180 Degrees Consulting Warwick CIC × Enactus Warwick Consulting Case Competition (University of Warwick) "
            "Developed a profitable, exit-ready growth and cost-optimisation strategy for a Series C $75M CCUS company, anchoring expansion in the Netherlands"
        )
        b = bullet()
        b.add_run(
            "3rd Place – NVIDIA ICETCI Hackathon (2023) – Out of 30 teams across India; "
            "published a research paper in ICETCI 2023 on data extraction, "
            "pre-processing and ensemble training of LLMs"
        )

        b = bullet()
    else:
        for x in ach_points:
            b = bullet()
            b.add_run(x)

    
    # ---------- Save ----------
    doc.save(file_name)


def create_cover_letter_docx(cover_letter_text, file_name):
    doc = Document()

    # ---------- Global font ----------
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # ---------- Margins ----------
    section = doc.sections[0]
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

    # ---------- Cover Letter Content ----------
    paragraphs = cover_letter_text.split('\n')
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        pf = p.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # ---------- Save ----------
    doc.save(file_name)